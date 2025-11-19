"""
Document Content Extraction Service
Extracts text content from uploaded files for LLM processing
"""

import io
import logging
from typing import Dict, Any, Optional
import xml.etree.ElementTree as ET

log = logging.getLogger(__name__)

# Try to import PDF parser
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf
        from pypdf import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        log.warning("PDF parsing not available - install PyPDF2 or pypdf")

class DocumentExtractor:
    """Extract text content from various document formats"""
    
    def __init__(self):
        self.max_content_length = 100000  # 100KB max per file for LLM context
        
    async def extract_content(self, file_data: bytes, filename: str, mime_type: str = None) -> Dict[str, Any]:
        """
        Extract text content from a file
        
        Returns:
            {
                'text_content': str,  # Extracted text
                'metadata': dict,     # File metadata
                'extraction_method': str,
                'success': bool
            }
        """
        
        result = {
            'text_content': '',
            'metadata': {},
            'extraction_method': 'none',
            'success': False,
            'error': None
        }
        
        try:
            # Determine file type
            file_extension = filename.split('.')[-1].lower() if '.' in filename else ''
            
            if file_extension == 'pdf' or (mime_type and 'pdf' in mime_type):
                result = await self._extract_pdf(file_data)
            elif file_extension in ['xml', 'html', 'htm']:
                result = await self._extract_xml_html(file_data, filename)
            elif file_extension in ['txt', 'md', 'csv', 'json', 'log']:
                result = await self._extract_text(file_data)
            elif file_extension in ['doc', 'docx']:
                result = await self._extract_doc(file_data)
            else:
                # Try as text fallback
                try:
                    text = file_data.decode('utf-8', errors='ignore')
                    if len(text) > 50 and text.isprintable():
                        result = await self._extract_text(file_data)
                    else:
                        result['error'] = f"Unsupported file type: {file_extension}"
                except:
                    result['error'] = f"Cannot extract from: {file_extension}"
            
            # Truncate if too long
            if result['text_content'] and len(result['text_content']) > self.max_content_length:
                log.info(f"Truncating content from {len(result['text_content'])} to {self.max_content_length} chars")
                result['text_content'] = result['text_content'][:self.max_content_length] + "\n\n[...Content truncated for length...]"
                result['truncated'] = True
            
            return result
            
        except Exception as e:
            log.error(f"Content extraction failed for {filename}: {e}")
            return {
                'text_content': '',
                'metadata': {},
                'extraction_method': 'failed',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_pdf(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            return {
                'text_content': '',
                'metadata': {},
                'extraction_method': 'pdf_unavailable',
                'success': False,
                'error': 'PDF library not installed'
            }
        
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PdfReader(pdf_file)
            
            # Extract text from all pages
            text_content = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n=== Page {page_num + 1} of {page_count} ===\n{page_text}\n"
                except Exception as e:
                    log.debug(f"Could not extract text from page {page_num + 1}: {e}")
            
            # Extract metadata
            pdf_metadata = {}
            if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                try:
                    for key, value in pdf_reader.metadata.items():
                        if value:
                            clean_key = key.replace('/', '')
                            pdf_metadata[clean_key] = str(value)
                except:
                    pass
            
            return {
                'text_content': text_content.strip(),
                'metadata': {
                    'page_count': page_count,
                    'pdf_metadata': pdf_metadata,
                    'word_count': len(text_content.split())
                },
                'extraction_method': 'pdf',
                'success': True
            }
            
        except Exception as e:
            return {
                'text_content': '',
                'metadata': {},
                'extraction_method': 'pdf_failed',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_xml_html(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from XML/HTML"""
        try:
            # First try as XML
            try:
                root = ET.fromstring(file_data)
                # Extract all text nodes
                text_content = ' '.join(root.itertext())
                extraction_method = 'xml'
            except:
                # Fallback to text extraction
                text_content = file_data.decode('utf-8', errors='ignore')
                # Remove HTML tags (basic)
                import re
                text_content = re.sub(r'<[^>]+>', ' ', text_content)
                text_content = re.sub(r'\s+', ' ', text_content)
                extraction_method = 'html_text'
            
            return {
                'text_content': text_content.strip(),
                'metadata': {
                    'file_type': 'xml/html',
                    'word_count': len(text_content.split())
                },
                'extraction_method': extraction_method,
                'success': True
            }
            
        except Exception as e:
            return {
                'text_content': '',
                'metadata': {},
                'extraction_method': 'xml_failed',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_text(self, file_data: bytes) -> Dict[str, Any]:
        """Extract plain text"""
        try:
            # Try multiple encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text_content = file_data.decode(encoding)
                    break
                except:
                    continue
            else:
                text_content = file_data.decode('utf-8', errors='ignore')
            
            return {
                'text_content': text_content.strip(),
                'metadata': {
                    'word_count': len(text_content.split()),
                    'line_count': len(text_content.split('\n'))
                },
                'extraction_method': 'text',
                'success': True
            }
            
        except Exception as e:
            return {
                'text_content': '',
                'metadata': {},
                'extraction_method': 'text_failed',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_doc(self, file_data: bytes) -> Dict[str, Any]:
        """Extract from DOC/DOCX (placeholder)"""
        try:
            # Try python-docx if available
            try:
                import docx
                doc_file = io.BytesIO(file_data)
                doc = docx.Document(doc_file)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                return {
                    'text_content': text_content.strip(),
                    'metadata': {
                        'paragraph_count': len(doc.paragraphs),
                        'word_count': len(text_content.split())
                    },
                    'extraction_method': 'docx',
                    'success': True
                }
            except ImportError:
                return {
                    'text_content': '',
                    'metadata': {},
                    'extraction_method': 'docx_unavailable',
                    'success': False,
                    'error': 'python-docx not installed'
                }
        except Exception as e:
            return {
                'text_content': '',
                'metadata': {},
                'extraction_method': 'doc_failed',
                'success': False,
                'error': str(e)
            }

# Global instance
document_extractor = DocumentExtractor()

